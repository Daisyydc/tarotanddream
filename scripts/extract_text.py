from __future__ import annotations

import json
import re
import traceback
from pathlib import Path
from typing import Dict, Any, List

try:
    import fitz  # pymupdf
except ImportError:
    raise ImportError("PyMuPDF is not installed. Please run: python -m pip install pymupdf")

from docx import Document
from tqdm import tqdm


RAW_PDF_DIR = Path("data/raw/pdf")
RAW_DOCX_DIR = Path("data/raw/docx")
OUTPUT_DIR = Path("data/processed/extracted")
REPORT_DIR = Path("data/processed/reports")

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
REPORT_DIR.mkdir(parents=True, exist_ok=True)


NOISE_PATTERNS = [
    r"http[s]?://\S+",
    r"www\.\S+",
    r"微信[:：]?\s*\S+",
    r"QQ[:：]?\s*\d+",
    r"扫码.*",
    r"二維碼.*",
    r"二维码.*",
    r"strc\.taobao\.com",
    r"上万套周易教程请加微信[:：]?\s*\S+",
    r"也可直接进入\s*\S+",
    r"猫鼬魔法仪式小铺",
    r"图片取自http\S+",
]


def count_chinese_chars(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def count_garbled_chars(text: str) -> int:
    return len(re.findall(r"[�◇◆□■△▽▼▲�\u0000-\u001f]", text))


def remove_noise_for_eval(text: str) -> str:
    cleaned = text
    for pattern in NOISE_PATTERNS:
        cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def assess_text_quality(texts: List[str]) -> Dict[str, Any]:
    merged = "\n".join(t for t in texts if t and t.strip())
    merged_eval = remove_noise_for_eval(merged)

    total_chars = len(merged_eval)
    chinese_chars = count_chinese_chars(merged_eval)
    garbled_chars = count_garbled_chars(merged_eval)

    chinese_ratio = chinese_chars / total_chars if total_chars > 0 else 0.0
    garbled_ratio = garbled_chars / total_chars if total_chars > 0 else 0.0
    nonempty_units = sum(1 for t in texts if t and t.strip())

    return {
        "total_chars": total_chars,
        "chinese_chars": chinese_chars,
        "garbled_chars": garbled_chars,
        "chinese_ratio": round(chinese_ratio, 4),
        "garbled_ratio": round(garbled_ratio, 4),
        "nonempty_units": nonempty_units,
    }


def decide_acceptance(file_type: str, quality_stats: Dict[str, Any]) -> Dict[str, Any]:
    total_chars = quality_stats["total_chars"]
    chinese_ratio = quality_stats["chinese_ratio"]
    garbled_ratio = quality_stats["garbled_ratio"]
    nonempty_units = quality_stats["nonempty_units"]

    accepted = False
    quality = "bad"
    reject_reason = None

    if file_type == "pdf":
        if total_chars >= 3000 and chinese_ratio >= 0.25 and garbled_ratio <= 0.05 and nonempty_units >= 10:
            accepted = True
            quality = "usable"
        elif total_chars >= 1200 and chinese_ratio >= 0.20 and garbled_ratio <= 0.10 and nonempty_units >= 5:
            accepted = False
            quality = "partial"
            reject_reason = "below_main_kb_threshold"
        else:
            accepted = False
            quality = "bad"
            reject_reason = "low_text_quality_or_insufficient_content"

    elif file_type == "docx":
        if total_chars >= 2000 and chinese_ratio >= 0.25 and garbled_ratio <= 0.03 and nonempty_units >= 20:
            accepted = True
            quality = "usable"
        elif total_chars >= 800 and chinese_ratio >= 0.20 and nonempty_units >= 8:
            accepted = False
            quality = "partial"
            reject_reason = "below_main_kb_threshold"
        else:
            accepted = False
            quality = "bad"
            reject_reason = "low_text_quality_or_insufficient_content"

    return {
        "accepted": accepted,
        "quality": quality,
        "reject_reason": reject_reason,
    }


def extract_pdf_text(pdf_path: Path) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "source_file": pdf_path.name,
        "file_type": "pdf",
        "book_id": pdf_path.stem,
        "pages": [],
        "stats": {
            "total_pages": 0,
            "success_pages": 0,
            "failed_pages": 0,
        },
    }

    doc = fitz.open(pdf_path)
    result["stats"]["total_pages"] = len(doc)

    for i, page in enumerate(doc, start=1):
        try:
            text = page.get_text("text")
            result["pages"].append({
                "page_num": i,
                "text": text.strip() if text else ""
            })
            result["stats"]["success_pages"] += 1
        except Exception as e:
            result["pages"].append({
                "page_num": i,
                "text": "",
                "error": str(e)
            })
            result["stats"]["failed_pages"] += 1

    page_texts = [p.get("text", "") for p in result["pages"]]
    quality_stats = assess_text_quality(page_texts)
    decision = decide_acceptance("pdf", quality_stats)

    result["quality_stats"] = quality_stats
    result.update(decision)
    return result


def extract_docx_text(docx_path: Path) -> Dict[str, Any]:
    doc = Document(docx_path)
    paragraphs: List[Dict[str, Any]] = []

    for i, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "para_num": i,
                "text": text
            })

    result = {
        "source_file": docx_path.name,
        "file_type": "docx",
        "book_id": docx_path.stem,
        "paragraphs": paragraphs,
        "stats": {
            "total_paragraphs": len(doc.paragraphs),
            "nonempty_paragraphs": len(paragraphs)
        }
    }

    para_texts = [p.get("text", "") for p in paragraphs]
    quality_stats = assess_text_quality(para_texts)
    decision = decide_acceptance("docx", quality_stats)

    result["quality_stats"] = quality_stats
    result.update(decision)
    return result


def save_json(data: Dict[str, Any], out_path: Path) -> None:
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def main() -> None:
    pdf_files = list(RAW_PDF_DIR.glob("*.pdf"))
    docx_files = list(RAW_DOCX_DIR.glob("*.docx"))

    print(f"Found {len(pdf_files)} PDF files.")
    print(f"Found {len(docx_files)} DOC/DOCX files.")

    report = {
        "pdf": [],
        "docx": [],
        "errors": []
    }

    for pdf_path in tqdm(pdf_files, desc="Extracting PDFs"):
        try:
            data = extract_pdf_text(pdf_path)
            out_path = OUTPUT_DIR / f"{pdf_path.stem}.json"
            save_json(data, out_path)

            report["pdf"].append({
                "source_file": pdf_path.name,
                "accepted": data["accepted"],
                "quality": data["quality"],
                "reject_reason": data["reject_reason"],
                "total_pages": data["stats"]["total_pages"],
                "success_pages": data["stats"]["success_pages"],
                "failed_pages": data["stats"]["failed_pages"],
                "quality_stats": data["quality_stats"],
            })
        except Exception as e:
            report["errors"].append({
                "source_file": pdf_path.name,
                "file_type": "pdf",
                "error": str(e),
                "traceback": traceback.format_exc()
            })
            print(f"[PDF ERROR] {pdf_path.name}: {e}")

    for docx_path in tqdm(docx_files, desc="Extracting DOCX files"):
        try:
            data = extract_docx_text(docx_path)
            out_path = OUTPUT_DIR / f"{docx_path.stem}.json"
            save_json(data, out_path)

            report["docx"].append({
                "source_file": docx_path.name,
                "accepted": data["accepted"],
                "quality": data["quality"],
                "reject_reason": data["reject_reason"],
                "total_paragraphs": data["stats"]["total_paragraphs"],
                "nonempty_paragraphs": data["stats"]["nonempty_paragraphs"],
                "quality_stats": data["quality_stats"],
            })
        except Exception as e:
            report["errors"].append({
                "source_file": docx_path.name,
                "file_type": "docx",
                "error": str(e),
                "traceback": traceback.format_exc()
            })
            print(f"[DOCX ERROR] {docx_path.name}: {e}")

    report_path = REPORT_DIR / "extraction_report.json"
    save_json(report, report_path)
    print(f"\nSaved report to: {report_path}")


if __name__ == "__main__":
    main()